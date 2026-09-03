"""Ingestion fiches terrain (.xlsx, .csv, .docx, .doc) vers lignes norme."""

from __future__ import annotations

import csv
import io
import shutil
import subprocess
import tempfile
from datetime import date
from pathlib import Path

from django.conf import settings

from apps.reports.norme import (
    ImportValidationError,
    NORME_COLUMNS,
    _align_values_to_headers,
    _friendly_error,
    _is_header_row,
    _is_payload_row,
    _is_section_separator,
    canonicalize_header,
    extract_period_from_text,
    normalize_row_keys,
    rows_from_csv,
    rows_from_xlsx,
    _to_float,
)


def _row_has_period(row: dict) -> bool:
    return bool(row.get('date_debut') not in (None, '') and row.get('date_fin') not in (None, ''))


def apply_period_to_rows(rows: list[dict], period: tuple[date, date] | None) -> list[dict]:
    if not period:
        return rows
    debut, fin = period
    for row in rows:
        if row.get('date_debut') in (None, ''):
            row['date_debut'] = debut
        if row.get('date_fin') in (None, ''):
            row['date_fin'] = fin
    return rows


def _convert_doc_to_docx(raw: bytes) -> bytes:
    soffice = shutil.which('soffice') or shutil.which('libreoffice')
    if not soffice:
        raise ImportValidationError(
            'Les fichiers Word .doc ne peuvent pas être lus sur ce serveur.',
            [
                _friendly_error(
                    row=None,
                    column=None,
                    message='Conversion .doc indisponible (LibreOffice absent côté serveur).',
                    how_to_fix='Enregistrez la fiche en .docx ou .xlsx depuis Word, puis déposez-la.',
                )
            ],
        )
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        src = tmp_path / 'fiche.doc'
        src.write_bytes(raw)
        profile = tmp_path / 'lo-profile'
        try:
            subprocess.run(
                [
                    soffice,
                    '--headless',
                    '--norestore',
                    f'-env:UserInstallation=file://{profile}',
                    '--convert-to',
                    'docx',
                    '--outdir',
                    str(tmp_path),
                    str(src),
                ],
                check=True,
                capture_output=True,
                timeout=90,
            )
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired) as exc:
            raise ImportValidationError(
                'La fiche Word .doc n’a pas pu être convertie.',
                [
                    _friendly_error(
                        row=None,
                        column=None,
                        message='La conversion vers .docx a échoué.',
                        how_to_fix='Dans Word : Fichier → Enregistrer sous → .docx ou .xlsx, puis déposez ce fichier.',
                    )
                ],
            ) from exc
        produced = list(tmp_path.glob('*.docx'))
        if not produced:
            raise ImportValidationError(
                'La fiche Word n’a pas produit de .docx.',
                [
                    _friendly_error(
                        row=None,
                        column=None,
                        message='Aucun .docx généré après conversion.',
                        how_to_fix='Enregistrez la fiche en .docx depuis Microsoft Word, puis déposez-la.',
                    )
                ],
            )
        return produced[0].read_bytes()


def _table_to_matrix(table) -> list[list[str]]:
    return [
        [(cell.text or '').replace('\n', ' ').strip() for cell in row.cells]
        for row in table.rows
    ]


def _matrix_to_rows(matrix: list[list[str]], extra_text: str = '') -> list[dict]:
    if not matrix:
        return []
    headers: list[str] = []
    header_idx = None
    for i, values in enumerate(matrix):
        if not values or all(not str(v).strip() for v in values):
            continue
        if _is_section_separator(values):
            continue
        if _is_header_row(values):
            headers = [canonicalize_header(v) for v in values]
            header_idx = i
            break
    if header_idx is None or not headers:
        return []
    rows: list[dict] = []
    for values in matrix[header_idx + 1 :]:
        if not values or all(not str(v).strip() for v in values):
            continue
        if _is_section_separator(values) or _is_header_row(values):
            if _is_header_row(values):
                headers = [canonicalize_header(v) for v in values]
            continue
        aligned = _align_values_to_headers(headers, values)
        raw = {key: aligned[idx] if idx < len(aligned) else None for idx, key in enumerate(headers) if key}
        normalized = normalize_row_keys(raw)
        if _is_payload_row(normalized):
            rows.append(normalized)
    period = extract_period_from_text(extra_text)
    if period:
        apply_period_to_rows(rows, period)
    return rows


def rows_from_docx(file_bytes: bytes, filename: str = '') -> list[dict]:
    try:
        from docx import Document
    except ImportError as exc:
        raise ImportValidationError(
            'Lecture des fichiers Word .docx indisponible.',
            [
                _friendly_error(
                    row=None,
                    column=None,
                    message='Le module python-docx n’est pas installé.',
                    how_to_fix='Enregistrez la fiche en .xlsx, ou installez python-docx sur le serveur.',
                )
            ],
        ) from exc
    try:
        document = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        raise ImportValidationError(
            'Le fichier Word n’a pas pu être ouvert.',
            [
                _friendly_error(
                    row=None,
                    column=None,
                    message='Le .docx est peut-être endommagé.',
                    how_to_fix='Ouvrez-le dans Word et enregistrez-le en .xlsx, puis réessayez.',
                )
            ],
        ) from exc
    paragraphs = ' '.join(p.text for p in document.paragraphs if p.text)
    title_text = f'{filename} {paragraphs}'
    best: list[dict] = []
    for table in document.tables:
        rows = _matrix_to_rows(_table_to_matrix(table), extra_text=title_text)
        if len(rows) > len(best):
            best = rows
    if not best:
        raise ImportValidationError(
            'Aucun tableau de relevés trouvé dans le Word.',
            [
                _friendly_error(
                    row=None,
                    column=None,
                    message='La fiche ne contient pas de tableau lisible (site, groupe, quantités).',
                    how_to_fix='Enregistrez la fiche en Excel (.xlsx), puis déposez-la.',
                )
            ],
        )
    period = extract_period_from_text(title_text)
    if period:
        apply_period_to_rows(best, period)
    return best


def load_rapport_rows_from_bytes(filename: str, raw: bytes) -> list[dict]:
    lower = (filename or 'rapport').lower()
    if lower.endswith('.xlsx'):
        rows = rows_from_xlsx(raw)
    elif lower.endswith('.csv'):
        rows = rows_from_csv(raw)
    elif lower.endswith('.docx'):
        rows = rows_from_docx(raw, filename=filename)
    elif lower.endswith('.doc'):
        rows = rows_from_docx(_convert_doc_to_docx(raw), filename=filename)
    else:
        raise ImportValidationError(
            'Type de fichier non accepté.',
            [
                _friendly_error(
                    row=None,
                    column=None,
                    message=f'Le fichier « {filename} » n’est pas un Excel, un CSV ou un Word.',
                    how_to_fix='Déposez un fichier .xlsx, .csv, .docx ou .doc.',
                )
            ],
        )
    period = extract_period_from_text(filename)
    if period:
        apply_period_to_rows(rows, period)
    if rows and not any(_row_has_period(r) for r in rows):
        blob = ' '.join(str(v) for row in rows[:3] for v in row.values() if v not in (None, ''))
        fallback = extract_period_from_text(filename, blob)
        if fallback:
            apply_period_to_rows(rows, fallback)
    return rows


def load_rapport_rows(path: str | Path) -> list[dict]:
    path = Path(path)
    if not path.exists():
        raise ImportValidationError(
            f'Fichier introuvable : {path}',
            [
                _friendly_error(
                    row=None,
                    column=None,
                    message=f'Le fichier « {path} » n’existe pas.',
                    how_to_fix='Vérifiez le chemin (.xlsx, .csv, .docx ou .doc).',
                )
            ],
        )
    return load_rapport_rows_from_bytes(path.name, path.read_bytes())


def write_preprocessed_export(rows: list[dict], date_debut, date_fin) -> Path | None:
    debut = date_debut.isoformat() if hasattr(date_debut, 'isoformat') else str(date_debut)
    fin = date_fin.isoformat() if hasattr(date_fin, 'isoformat') else str(date_fin)
    out_dir = Path(settings.PROJECT_ROOT) / 'data' / 'exports' / 'rapports' / f'{debut}_{fin}'
    try:
        out_dir.mkdir(parents=True, exist_ok=True)
        dest = out_dir / 'lignes.csv'
        buffer = io.StringIO()
        writer = csv.DictWriter(buffer, fieldnames=NORME_COLUMNS, extrasaction='ignore', lineterminator='\n')
        writer.writeheader()
        for row in rows:
            payload = {col: row.get(col, '') for col in NORME_COLUMNS}
            for key in ('date_debut', 'date_fin'):
                value = payload.get(key)
                if hasattr(value, 'isoformat'):
                    payload[key] = value.isoformat()
            for key in (
                'quantités_cuve_principale',
                'quantite_cuve_journaliere',
                'depotage',
                'compteur_horaire',
            ):
                payload[key] = _to_float(payload.get(key))
            writer.writerow(payload)
        dest.write_text(buffer.getvalue(), encoding='utf-8-sig')
        return dest
    except OSError:
        return None
